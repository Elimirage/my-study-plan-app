import io
import json
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from ai import call_yandex_lite


def _safe_str(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        if value is None or value == "":
            return default
        return int(float(value))
    except Exception:
        return default


def _normalize_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [_safe_str(x) for x in value if _safe_str(x)]
    if isinstance(value, str):
        value = value.strip()
        if not value:
            return []
        if value.startswith("[") and value.endswith("]"):
            try:
                parsed = json.loads(value)
                if isinstance(parsed, list):
                    return [_safe_str(x) for x in parsed if _safe_str(x)]
            except Exception:
                pass
        return [x.strip() for x in value.split(",") if x.strip()]
    return [_safe_str(value)]


def _safe_json_from_text(text: str) -> Dict[str, Any]:
    start = text.find("{")
    end = text.rfind("}") + 1
    if start == -1 or end <= 0:
        raise ValueError(f"Не найден JSON в ответе модели:\n{text[:1500]}")
    return json.loads(text[start:end])


def _extract_competencies_only(row: Dict[str, Any]) -> List[str]:
    raw = row.get("Компетенции ФГОС", [])
    items = _normalize_list(raw)

    result = []
    for item in items:
        upper = item.upper().replace(" ", "")
        if upper.startswith("УК-") or upper.startswith("ОПК-") or upper.startswith("ПК-"):
            result.append(item)

    seen = set()
    unique = []
    for item in result:
        if item not in seen:
            unique.append(item)
            seen.add(item)
    return unique


def _extract_tf_only(row: Dict[str, Any]) -> List[str]:
    raw = row.get("Трудовые функции", [])
    items = _normalize_list(raw)

    result = []
    for item in items:
        cleaned = _safe_str(item)
        if "/" in cleaned or cleaned[:1].isalpha():
            result.append(cleaned)

    seen = set()
    unique = []
    for item in result:
        if item not in seen:
            unique.append(item)
            seen.add(item)
    return unique


def _set_base_style(doc: Document) -> None:
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)


def _set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    size: int = 11,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(_safe_str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size: int = 12,
    space_after: int = 6,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(_safe_str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)


def _add_signature_line(doc: Document, title: str, person: str = "") -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    row = table.rows[0].cells
    _set_cell_text(row[0], title, size=12)
    _set_cell_text(row[1], "__________________", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    _set_cell_text(
        row[2],
        person if person else "__________________",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
    )


def generate_work_program_content(
    discipline_row: Dict[str, Any],
    profile: str,
    direction_code: str,
    direction_name: str,
    qualification: str,
    education_form: str,
    university_name: str,
    faculty_name: str,
) -> Dict[str, Any]:
    discipline_name = _safe_str(discipline_row.get("Дисциплина"))
    semester = _safe_int(discipline_row.get("Семестр"), 1)
    hours = _safe_int(discipline_row.get("Часы"), 108)
    control_form = _safe_str(discipline_row.get("Форма контроля")) or "зачёт"

    competencies = _extract_competencies_only(discipline_row)
    tf_list = _extract_tf_only(discipline_row)
    reason = _safe_str(discipline_row.get("Обоснование"))

    credits = max(1, round(hours / 36)) if hours else 3

    competencies_text = ", ".join(competencies) if competencies else "не указаны"
    tf_text = ", ".join(tf_list) if tf_list else "не указаны"

    prompt = f"""
Ты — опытный методист российского вуза.

Сгенерируй рабочую программу дисциплины в формате СТРОГОГО JSON.
Без markdown.
Без пояснений.
Без текста до JSON и после JSON.

Входные данные:
- Дисциплина: {discipline_name}
- Профиль: {profile}
- Код направления: {direction_code}
- Направление подготовки: {direction_name}
- Квалификация: {qualification}
- Форма обучения: {education_form}
- Семестр: {semester}
- Объем часов: {hours}
- Зачетные единицы: {credits}
- Форма контроля: {control_form}
- Компетенции ФГОС: {competencies_text}
- Трудовые функции профстандарта: {tf_text}
- Обоснование дисциплины: {reason if reason else "не указано"}

СТРОГИЕ ПРАВИЛА:
1. Используй только компетенции ФГОС (УК-..., ОПК-..., ПК-...) в таблице результатов.
2. Не включай трудовые функции профстандарта в таблицу результатов.
3. Трудовые функции можно упоминать только в описательном тексте, если это уместно.
4. Содержание должно соответствовать дисциплине "{discipline_name}".
5. Сделай документ реалистичным для вуза РФ.
6. Минимум:
   - 8 тем лекций,
   - 4 лабораторных или практических работ,
   - 4 строки самостоятельной работы,
   - 5 источников литературы.
7. Таблица структуры должна содержать не менее 6 строк.
8. Самостоятельная работа должна быть разнообразной:
   - подготовка к занятиям,
   - анализ практики,
   - реферат,
   - решение кейсов.
9. Не используй "..." и пустые заглушки.
10. Программное обеспечение и оборудование должны быть конкретными.
11. Используй минимум 1 УК, 1 ОПК и 1 ПК.
12. Если дисциплина гуманитарная, не вставляй ИТ-содержание.
13. Если дисциплина техническая, не вставляй юридическое содержание.

Верни JSON такой структуры:
{{
  "title": "Рабочая программа дисциплины",
  "discipline_code": "Б1.О.01",
  "discipline_name": "{discipline_name}",
  "goals": "1-2 абзаца",
  "place_in_program": "1-2 абзаца",
  "results": [
    {{
      "code": "УК-1",
      "competence": "Формулировка компетенции",
      "indicator": "Индикатор достижения",
      "know": "Что знает",
      "able": "Что умеет",
      "master": "Чем владеет"
    }}
  ],
  "total_credits": {credits},
  "total_hours": {hours},
  "structure_rows": [
    {{
      "section": "Раздел 1. ...",
      "topic": "Тема 1.1. ...",
      "semester": {semester},
      "weeks": "1-2",
      "lectures": 2,
      "labs": 2,
      "other_contact": 0,
      "self_study": 6,
      "current_control": "Опрос",
      "intermediate_control": "{control_form}"
    }}
  ],
  "lecture_topics": [
    "Тема 1.1. ..."
  ],
  "lab_topics": [
    {{
      "name": "Практическое занятие 1. ...",
      "hours": 2
    }}
  ],
  "education_technologies": [
    "Проблемное обучение"
  ],
  "self_study_rows": [
    {{
      "weeks": "1-2",
      "topic": "Тема 1.1. ...",
      "kind": "Подготовка к занятиям",
      "task": "Конкретное задание",
      "literature": "1-3",
      "hours": 6
    }}
  ],
  "assessment_tools": [
    "Опрос",
    "Проверка практических работ",
    "{control_form}"
  ],
  "literature": [
    "1. ..."
  ],
  "software": [
    "КонсультантПлюс"
  ],
  "equipment": [
    "Компьютерный класс"
  ]
}}
"""

    raw = call_yandex_lite(
        [{"role": "user", "text": prompt}],
        temperature=0.2,
        max_tokens=3500,
    )

    data = _safe_json_from_text(raw)

    data["meta"] = {
        "university_name": university_name,
        "faculty_name": faculty_name,
        "direction_code": direction_code,
        "direction_name": direction_name,
        "profile": profile,
        "qualification": qualification,
        "education_form": education_form,
        "input_competencies": competencies,
        "input_tf": tf_list,
    }

    return data


def create_work_program_docx(data: Dict[str, Any]) -> bytes:
    doc = Document()
    _set_base_style(doc)

    meta = data.get("meta", {})
    university_name = _safe_str(meta.get("university_name"))
    faculty_name = _safe_str(meta.get("faculty_name"))
    direction_code = _safe_str(meta.get("direction_code"))
    direction_name = _safe_str(meta.get("direction_name"))
    profile = _safe_str(meta.get("profile"))
    qualification = _safe_str(meta.get("qualification"))
    education_form = _safe_str(meta.get("education_form"))

    discipline_code = _safe_str(data.get("discipline_code")) or "Б1.О.01"
    discipline_name = _safe_str(data.get("discipline_name"))
    total_credits = _safe_int(data.get("total_credits"), 3)
    total_hours = _safe_int(data.get("total_hours"), 108)

    current_year = datetime.now().year
    city_name = "Пенза"

    # Блок "Утверждаю"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("УТВЕРЖДАЮ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Декан факультета")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("__________________ /__________________/")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'"____" __________ {current_year} г.')
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    _add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Титульный лист
    _add_paragraph(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "РОССИЙСКОЙ ФЕДЕРАЦИИ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, university_name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, faculty_name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "РАБОЧАЯ ПРОГРАММА ДИСЦИПЛИНЫ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"{discipline_code} {discipline_name}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    _add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_paragraph(doc, f"Направление подготовки: {direction_code} «{direction_name}»", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, f"Направленность подготовки: «{profile}»", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, f"Квалификация выпускника: {qualification}", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, f"Форма обучения: {education_form}", align=WD_ALIGN_PARAGRAPH.LEFT)

    _add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"{city_name}, {current_year}", align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # 1
    _add_heading(doc, "1. Цели освоения дисциплины")
    _add_paragraph(doc, _safe_str(data.get("goals")))

    # 2
    _add_heading(doc, "2. Место дисциплины в структуре ОПОП")
    _add_paragraph(doc, _safe_str(data.get("place_in_program")))

    # 3
    _add_heading(doc, f"3. Результаты освоения дисциплины «{discipline_name}»")

    results = data.get("results", [])
    if isinstance(results, list) and results:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        _set_cell_text(hdr[0], "Код", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(hdr[1], "Компетенция", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(hdr[2], "Индикатор", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(hdr[3], "Знать", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(hdr[4], "Уметь", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(hdr[5], "Владеть", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        for item in results:
            code = _safe_str(item.get("code")).upper().replace(" ", "")
            if not (code.startswith("УК-") or code.startswith("ОПК-") or code.startswith("ПК-")):
                continue

            row = table.add_row().cells
            _set_cell_text(row[0], item.get("code", ""))
            _set_cell_text(row[1], item.get("competence", ""))
            _set_cell_text(row[2], item.get("indicator", ""))
            _set_cell_text(row[3], item.get("know", ""))
            _set_cell_text(row[4], item.get("able", ""))
            _set_cell_text(row[5], item.get("master", ""))

    # 4
    _add_heading(doc, f"4. Структура и содержание дисциплины «{discipline_name}»")
    _add_paragraph(
        doc,
        f"Общая трудоемкость дисциплины составляет {total_credits} зачетных единиц, {total_hours} часов."
    )

    structure_rows = data.get("structure_rows", [])
    if isinstance(structure_rows, list) and structure_rows:
        table = doc.add_table(rows=1, cols=10)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        headers = [
            "Раздел",
            "Тема",
            "Семестр",
            "Недели",
            "Лекции",
            "Лаб./практ.",
            "Иное",
            "СРС",
            "Текущий контроль",
            "Промежуточный контроль",
        ]
        for i, h in enumerate(headers):
            _set_cell_text(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        for item in structure_rows:
            row = table.add_row().cells
            _set_cell_text(row[0], item.get("section", ""))
            _set_cell_text(row[1], item.get("topic", ""))
            _set_cell_text(row[2], item.get("semester", ""))
            _set_cell_text(row[3], item.get("weeks", ""))
            _set_cell_text(row[4], item.get("lectures", ""))
            _set_cell_text(row[5], item.get("labs", ""))
            _set_cell_text(row[6], item.get("other_contact", ""))
            _set_cell_text(row[7], item.get("self_study", ""))
            _set_cell_text(row[8], item.get("current_control", ""))
            _set_cell_text(row[9], item.get("intermediate_control", ""))

    _add_heading(doc, "4.2.1. Содержание лекционных занятий")
    for topic in data.get("lecture_topics", []):
        _add_paragraph(doc, f"• {_safe_str(topic)}")

    _add_heading(doc, "4.2.2. Темы лабораторных работ")
    lab_topics = data.get("lab_topics", [])
    if isinstance(lab_topics, list) and lab_topics:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        _set_cell_text(
            hdr[0],
            "Наименование лабораторной / практической работы",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER
        )
        _set_cell_text(hdr[1], "Часы", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        for item in lab_topics:
            row = table.add_row().cells
            _set_cell_text(row[0], item.get("name", ""))
            _set_cell_text(row[1], item.get("hours", ""))

    # 5
    _add_heading(doc, "5. Образовательные технологии")
    for item in data.get("education_technologies", []):
        _add_paragraph(doc, f"• {_safe_str(item)}")

    # 6
    _add_heading(
        doc,
        "6. Учебно-методическое обеспечение самостоятельной работы студентов. "
        "Оценочные средства для текущего контроля успеваемости и промежуточной аттестации."
    )

    ss_rows = data.get("self_study_rows", [])
    if isinstance(ss_rows, list) and ss_rows:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        headers = ["Недели", "Тема", "Вид работы", "Задание", "Литература", "Часы"]
        for i, h in enumerate(headers):
            _set_cell_text(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        for item in ss_rows:
            row = table.add_row().cells
            _set_cell_text(row[0], item.get("weeks", ""))
            _set_cell_text(row[1], item.get("topic", ""))
            _set_cell_text(row[2], item.get("kind", ""))
            _set_cell_text(row[3], item.get("task", ""))
            _set_cell_text(row[4], item.get("literature", ""))
            _set_cell_text(row[5], item.get("hours", ""))

    _add_heading(doc, "6.3. Оценочные средства")
    for item in data.get("assessment_tools", []):
        _add_paragraph(doc, f"• {_safe_str(item)}")

    # 7
    _add_heading(doc, "7. Учебно-методическое и материально-техническое обеспечение дисциплины")

    _add_paragraph(doc, "а) Литература", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    for item in data.get("literature", []):
        _add_paragraph(doc, _safe_str(item))

    _add_paragraph(doc, "б) Программное обеспечение", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    for item in data.get("software", []):
        _add_paragraph(doc, f"• {_safe_str(item)}")

    _add_paragraph(doc, "в) Материально-техническое обеспечение", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    for item in data.get("equipment", []):
        _add_paragraph(doc, f"• {_safe_str(item)}")

    # Подписи в конце
    _add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_heading(doc, "8. Лист согласования")
    _add_signature_line(doc, "Разработчик рабочей программы")
    _add_signature_line(doc, "Заведующий кафедрой")
    _add_signature_line(doc, "Председатель методической комиссии")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()